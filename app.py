import os
import uuid
import threading
import fitz  # PyMuPDF
import traceback
import io
from PIL import Image
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# --- Basic App Setup ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": ["https://jkmpdf.github.io", "https://www.jkmedit.in", "https://jkmedit.in"]}})

# --- Configuration ---
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

jobs = {}

# === THE NEW DOCUMENT RECONSTRUCTION ENGINE ===
def convert_with_document_reconstruction(pdf_path, docx_path):
    try:
        pdf_document = fitz.open(pdf_path)
        word_document = Document()
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            page_width = page.rect.width
            
            # --- STEP 1: GATHER ALL ELEMENTS ---
            # Get text blocks and their bounding boxes
            text_blocks = page.get_text("dict", flags=0)["blocks"]
            
            # Get drawings (vector graphics like signatures) and images
            # We will render them to PNGs to insert them
            all_elements = []

            # Add text elements
            for block in text_blocks:
                if 'lines' in block:
                    all_elements.append({
                        "type": "text",
                        "bbox": block["bbox"],
                        "content": block
                    })
            
            # Add image elements (both raster and vector)
            # This is more robust for finding things like signatures
            for drawing in page.get_drawings():
                # For each drawing, get its bounding box and render it
                rect = drawing['rect']
                if rect.is_empty or rect.width < 1 or rect.height < 1:
                    continue
                pix = page.get_pixmap(clip=rect, dpi=200)
                all_elements.append({
                    "type": "image",
                    "bbox": (rect.x0, rect.y0, rect.x1, rect.y1),
                    "content": pix.tobytes()
                })

            # --- STEP 2: SORT ALL ELEMENTS BY VERTICAL POSITION ---
            # This is the key to preserving the document's flow
            all_elements.sort(key=lambda item: item["bbox"][1]) # Sort by y0 (top coordinate)

            # --- STEP 3: REBUILD THE DOCUMENT IN THE CORRECT ORDER ---
            last_y1 = 0
            current_paragraph = None
            
            for element in all_elements:
                bbox = element["bbox"]
                
                if element["type"] == "text":
                    # --- Advanced Paragraph Logic ---
                    # Check vertical distance to decide if it's a new paragraph
                    if current_paragraph is None or bbox[1] > (last_y1 + 10): # 10 is a threshold
                        current_paragraph = word_document.add_paragraph()
                        # Set alignment for the new paragraph
                        if bbox[0] > page_width * 0.6:
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif (bbox[0] > page_width * 0.35 and bbox[2] < page_width * 0.65):
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add text to the current paragraph
                    for line in element["content"]["lines"]:
                        line_text = "".join([span["text"] for span in line["spans"]])
                        current_paragraph.add_run(line_text + " ")
                    
                    last_y1 = bbox[3] # Update the last vertical position

                elif element["type"] == "image":
                    # We are done with the previous paragraph
                    current_paragraph = None 
                    try:
                        image_stream = io.BytesIO(element["content"])
                        word_document.add_picture(image_stream)
                    except Exception as e:
                        print(f"Could not add image element: {e}")
                    last_y1 = bbox[3]
            
            if page_num < len(pdf_document) - 1:
                word_document.add_page_break()
                
        word_document.save(docx_path)
        return True, "Document Reconstruction successful"
    except Exception as e:
        print(f"!!! DOCUMENT RECONSTRUCTION FAILED. Error: {e}")
        traceback.print_exc()
        return False, str(e)

# --- Worker Function (now calls the new reconstruction engine) ---
def process_file(job_id, pdf_path, docx_path):
    print(f"--- Starting Document Reconstruction for job {job_id} ---")
    jobs[job_id]['status'] = 'PROCESSING'
    success, message = convert_with_document_reconstruction(pdf_path, docx_path)
    if success:
        jobs[job_id]['status'] = 'COMPLETED'
        print(f"--- Document Reconstruction COMPLETED for job {job_id} ---")
    else:
        jobs[job_id]['status'] = 'FAILED'
        jobs[job_id]['error'] = message
        print(f"--- Document Reconstruction FAILED for job {job_id} ---")

# --- API Endpoints (No changes needed here) ---
# ... (The rest of your app.py remains the same) ...

@app.route('/api/ocr/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files: return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '': return jsonify({"error": "No selected file"}), 400
    if file:
        job_id = str(uuid.uuid4())
        filename = f"{job_id}.pdf"
        output_filename = f"{job_id}.docx"
        pdf_path = os.path.join(UPLOAD_FOLDER, filename)
        docx_path = os.path.join(OUTPUT_FOLDER, output_filename)
        file.save(pdf_path)
        jobs[job_id] = {'status': 'PENDING', 'pdf_path': pdf_path, 'docx_path': docx_path}
        print(f"File saved to {pdf_path}. Starting thread for job {job_id}.")
        thread = threading.Thread(target=process_file, args=(job_id, pdf_path, docx_path))
        thread.start()
        return jsonify({"jobId": job_id})
    return jsonify({"error": "An unknown error occurred"}), 500

@app.route('/api/ocr/status/<job_id>', methods=['GET'])
def get_status(job_id):
    job = jobs.get(job_id)
    if not job: return jsonify({"error": "Job not found"}), 404
    response = {"status": job['status']}
    if job['status'] == 'FAILED': response['error'] = job.get('error', 'An unknown error occurred.')
    return jsonify(response)

@app.route('/api/ocr/download/<job_id>', methods=['GET'])
def download_file(job_id):
    job = jobs.get(job_id)
    if not job or job['status'] != 'COMPLETED': return jsonify({"error": "File not ready or job failed"}), 404
    try:
        return send_from_directory(OUTPUT_FOLDER, f"{job_id}.docx", as_attachment=True, download_name=f"converted_{job_id}.docx")
    except FileNotFoundError:
        return jsonify({"error": "File not found."}), 404

@app.route('/')
def index():
    return "JKM Edit Backend is running!"

if __name__ == '__main__':
    app.run(debug=True)
